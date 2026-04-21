"""
DOCX Utilities - Extract text from Word documents.

Educational Note: These utilities extract text from DOCX files using python-docx.
DOCX files are ZIP archives containing XML documents. The python-docx library
handles the complexity of parsing this structure.

We preserve basic document structure by converting to markdown-like format:
- Headings -> # Heading, ## Subheading, etc.
- Lists -> - item or 1. item
- Tables -> markdown table format
- Paragraphs -> plain text with blank lines between

The extracted text is then split into pages using the same logic as text files
(character count + natural break points).
"""
import logging
from pathlib import Path
from typing import Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX processing unavailable.")


def is_available() -> bool:
    """Check if python-docx library is available."""
    return DOCX_AVAILABLE


def extract_text_from_docx(docx_path: Path) -> Dict[str, Any]:
    """
    Extract text content from a DOCX file.

    Educational Note: This function reads the DOCX file and extracts:
    - All paragraphs (with heading level detection)
    - All tables (converted to markdown format)
    - Preserves document order

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Dict with:
            - success: bool
            - text: extracted text content
            - paragraph_count: number of paragraphs
            - table_count: number of tables
            - error: error message if failed
    """
    if not DOCX_AVAILABLE:
        return {
            "success": False,
            "error": "python-docx library not installed"
        }

    if not docx_path.exists():
        return {
            "success": False,
            "error": f"File not found: {docx_path}"
        }

    try:
        # Load the document
        doc = Document(docx_path)

        # Extract content in document order
        extracted_text = _extract_document_content(doc)

        # Count elements
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)

        return {
            "success": True,
            "text": extracted_text,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "extracted_at": datetime.now().isoformat()
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to extract text: {str(e)}"
        }


def _extract_document_content(doc: 'Document') -> str:
    """
    Extract all content from document in order.

    Educational Note: DOCX documents have a body that contains paragraphs
    and tables in order. We iterate through the body elements to maintain
    the correct document order.

    Args:
        doc: python-docx Document object

    Returns:
        Extracted text with markdown-like formatting
    """
    content_parts = []

    # Iterate through body elements in order
    # This maintains the correct order of paragraphs and tables
    for element in doc.element.body:
        # Check if it's a paragraph
        if element.tag.endswith('p'):
            para = Paragraph(element, doc)
            para_text = _extract_paragraph(para)
            if para_text:
                content_parts.append(para_text)

        # Check if it's a table
        elif element.tag.endswith('tbl'):
            table = Table(element, doc)
            table_text = _extract_table(table)
            if table_text:
                content_parts.append(table_text)

    # Join with double newlines for paragraph separation
    return '\n\n'.join(content_parts)


def _extract_paragraph(para: 'Paragraph') -> str:
    """
    Extract text from a paragraph with heading detection.

    Educational Note: Word documents have paragraph styles that indicate
    headings (Heading 1, Heading 2, etc.). We convert these to markdown
    format for better structure preservation.

    Args:
        para: python-docx Paragraph object

    Returns:
        Formatted paragraph text
    """
    text = para.text.strip()
    if not text:
        return ""

    # Check for heading style
    style_name = para.style.name if para.style else ""

    # Convert headings to markdown format
    if style_name.startswith('Heading'):
        try:
            # Extract heading level (Heading 1 -> #, Heading 2 -> ##, etc.)
            level = int(style_name.split()[-1])
            prefix = '#' * min(level, 6)  # Max 6 levels in markdown
            return f"{prefix} {text}"
        except (ValueError, IndexError):
            # If we can't parse the level, treat as regular paragraph
            pass

    # Check for list items
    if para.style and 'List' in para.style.name:
        # Detect if numbered or bulleted
        if 'Number' in para.style.name or 'Ordered' in para.style.name:
            return f"1. {text}"
        else:
            return f"- {text}"

    return text


def _extract_table(table: 'Table') -> str:
    """
    Extract table content in markdown format.

    Educational Note: Tables are converted to markdown table format:
    | Header 1 | Header 2 |
    |----------|----------|
    | Cell 1   | Cell 2   |

    This preserves the tabular structure in a text-friendly format.

    Args:
        table: python-docx Table object

    Returns:
        Markdown-formatted table string
    """
    if not table.rows:
        return ""

    rows_data = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            # Get cell text, replace newlines with spaces
            cell_text = cell.text.strip().replace('\n', ' ')
            row_cells.append(cell_text)
        rows_data.append(row_cells)

    if not rows_data:
        return ""

    # Build markdown table
    lines = []

    # First row as header
    header = rows_data[0]
    lines.append('| ' + ' | '.join(header) + ' |')

    # Separator line
    separator = '| ' + ' | '.join(['---'] * len(header)) + ' |'
    lines.append(separator)

    # Data rows
    for row in rows_data[1:]:
        # Ensure row has same number of cells as header
        while len(row) < len(header):
            row.append('')
        lines.append('| ' + ' | '.join(row[:len(header)]) + ' |')

    return '\n'.join(lines)
